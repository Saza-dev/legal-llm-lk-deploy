import streamlit as st
from docx import Document
from io import BytesIO
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain

from config.setup import retriever, llm
from utils.prompts import get_drafter_prompt

def run():
    st.title("Drafter")
    st.write("Describe the document you want to draft.")

    draft_input = st.text_input("Explain what kind of draft you need?")

    if draft_input:
        prompt = get_drafter_prompt()
        qa_chain = create_stuff_documents_chain(llm, prompt)
        rag_chain = create_retrieval_chain(retriever, qa_chain)

        response = rag_chain.invoke({"input": draft_input})
        st.write(response["answer"])

        doc = Document()
        doc.add_heading("Drafted Document", level=1)
        doc.add_paragraph(response["answer"])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📄 Download as Word (.docx)",
            data=buffer,
            file_name="drafted_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
